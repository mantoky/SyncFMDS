from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(r"C:\Users\robso\OneDrive\Documentos\New project")
IMG_PATH = BASE_DIR / "mop_agenda_fluxo.png"
OUT_DOCX = BASE_DIR / "MOP_Manual_Agenda_Inteligente_Infratech.docx"


def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    run = title.add_run("MOP + Manual do APP\nAgenda Inteligente Infratech")
    run.bold = True
    run.font.size = Pt(22)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph(
        f"Documento operacional para Microsoft Lists, Power Automate, Power Apps e Power BI\n"
        f"Data de emissão: {datetime.now():%d/%m/%Y}"
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_intro(doc: Document) -> None:
    doc.add_heading("1. Objetivo", level=1)
    doc.add_paragraph(
        "Padronizar a operação da agenda compartilhada do time Infratech, com fluxo de solicitação, "
        "detecção de conflitos de horário, validação das áreas responsáveis, aprovação final administrativa "
        "e governança de indicadores para acompanhamento em dashboard."
    )

    doc.add_heading("2. Visão da Solução", level=1)
    doc.add_paragraph(
        "A figura abaixo apresenta a aparência sugerida da solução e o fluxo completo entre as ferramentas."
    )
    if IMG_PATH.exists():
        doc.add_picture(str(IMG_PATH), width=Inches(6.6))
        cap = doc.add_paragraph("Figura 1 - Arquitetura visual (List + Automate + Power Apps + BI)")
        cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        doc.add_paragraph(
            "Imagem de referência não encontrada em: " + str(IMG_PATH)
        )


def add_list_config(doc: Document) -> None:
    doc.add_heading("3. Configuração Simulada - Microsoft Lists", level=1)
    doc.add_paragraph("Nome da lista: Agenda_Solicitacoes")

    headers = ["Campo", "Tipo", "Obrigatório", "Exemplo/Regra"]
    rows = [
        ("Solicitante", "Pessoa", "Sim", "Nome do responsável pela demanda"),
        ("EquipeInfratech", "Escolha", "Sim", "LTE SONDA, ALCON RÁDIO, CABLING SONDA, NETWORK SONDA, ENGEFAME, STEFANIFNI, STAFF VALE, PCM, SESMT"),
        ("TipoPedido", "Escolha", "Sim", "Emergencial, Programada, Horas de OMs"),
        ("DemandaAssociadaTipo", "Escolha", "Sim", "OS, OM, INCIDENTE, TASK, CHG, OUTROS"),
        ("DemandaAssociadaDetalhe", "Texto longo", "Não", "Ex.: CHG009876 - Janela backbone"),
        ("Inicio", "Data e hora", "Sim", "Data/hora inicial"),
        ("Fim", "Data e hora", "Sim", "Data/hora final"),
        ("AgendaEmConflito", "Sim/Não", "Sim", "Default: Não"),
        ("AssociarCom", "Escolha", "Sim", "Mesmo catálogo das equipes"),
        ("StatusPedido", "Escolha", "Sim", "Aguardando, Em validação, Aprovado, Reprovado, Rejeitado, Reagendamento"),
        ("ParecerAdministrador", "Texto longo", "Condicional", "Obrigatório em decisão final"),
        ("AprovadorFinal", "Pessoa", "Não", "Preenchido no fluxo"),
        ("DataDecisao", "Data e hora", "Não", "Preenchido no fluxo"),
        ("AcaoSolicitante", "Escolha", "Sim", "Rascunho ou Enviar para validação"),
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.add_paragraph("Validação da lista (List validation): [Inicio] < [Fim].")
    doc.add_paragraph("Views recomendadas: Calendário, Pendentes, Conflitos, Aprovados, Rejeitados.")


def add_automate_config(doc: Document) -> None:
    doc.add_heading("4. Configuração Simulada - Power Automate", level=1)

    doc.add_heading("Fluxo A - Triagem de Conflito", level=2)
    steps_a = [
        "Trigger: When an item is created or modified (SharePoint).",
        "Condição: processar somente quando AcaoSolicitante = Enviar para validação.",
        "Consultar itens do mesmo AssociarCom.",
        "Aplicar regra de sobreposição: Inicio_existente < Fim_novo E Fim_existente > Inicio_novo.",
        "Se houver conflito: AgendaEmConflito = Sim e StatusPedido = Em validação.",
        "Se não houver: AgendaEmConflito = Não e StatusPedido = Em validação.",
    ]
    for s in steps_a:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Fluxo B - Validação das Áreas Responsáveis", level=2)
    steps_b = [
        "Iniciar aprovação para STAFF, PCM, SESMT e GESTÃO SONDA.",
        "Se validado: encaminhar para aprovação final ADM.",
        "Se recusado: StatusPedido = Reprovado e parecer obrigatório.",
    ]
    for s in steps_b:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Fluxo C - Aprovação Final ADM", level=2)
    steps_c = [
        "Aprovado: StatusPedido = Aprovado (bolinha verde).",
        "Não aprovado: StatusPedido = Reprovado/Rejeitado/Reagendamento (bolinha vermelha).",
        "Registrar ParecerAdministrador, AprovadorFinal e DataDecisao.",
    ]
    for s in steps_c:
        doc.add_paragraph(s, style="List Number")

    doc.add_heading("Fluxo D - Bloqueio de Repetição após Rejeição", level=2)
    doc.add_paragraph(
        "Ao detectar nova solicitação no mesmo intervalo de um pedido rejeitado para o mesmo AssociarCom, "
        "forçar status Reagendamento e notificar o solicitante para escolher novo horário."
    )


def add_powerapps_config(doc: Document) -> None:
    doc.add_heading("5. Configuração Simulada - Power Apps", level=1)

    doc.add_heading("Tela 1 - Nova Solicitação (Solicitante)", level=2)
    for item in [
        "Formulário simplificado com os campos da lista.",
        "Botão Salvar rascunho: mantém AcaoSolicitante = Rascunho.",
        "Botão Enviar: define AcaoSolicitante = Enviar para validação.",
        "Validação de datas: Inicio deve ser menor que Fim.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Tela 2 - Gestão ADM", level=2)
    for item in [
        "Galeria com filtros por Status, Equipe, TipoPedido e período.",
        "Ações: Aprovar, Reprovar e Reagendar.",
        "ParecerAdministrador obrigatório para decisões não aprovadas.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Semáforo de Status", level=2)
    for item in [
        "Amarelo: Aguardando ou Em validação.",
        "Verde: Aprovado.",
        "Vermelho: Reprovado, Rejeitado ou Reagendamento.",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def add_user_manual(doc: Document) -> None:
    doc.add_heading("6. Manual de Uso do APP (Após Finalização)", level=1)

    doc.add_heading("Perfil Solicitante", level=2)
    for step in [
        "Acesse o APP Agenda Infratech.",
        "Clique em Nova Solicitação.",
        "Preencha equipe, tipo, demanda associada, início, fim e associar com.",
        "Clique em Enviar para validação.",
        "Acompanhe o status na tela Minhas Solicitações.",
        "Se o status for Rejeitado/Reagendamento, crie uma nova solicitação em horário diferente.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Perfil Validador (STAFF, PCM, SESMT, GESTÃO SONDA)", level=2)
    for step in [
        "Abra a fila Pendentes de validação.",
        "Analise conflito, impacto operacional e prioridade.",
        "Aprove ou recuse na etapa de validação.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Perfil Administrador", level=2)
    for step in [
        "Acesse a tela Gestão ADM.",
        "Filtre solicitações Em validação.",
        "Defina decisão final: Aprovar, Reprovar ou Reagendar.",
        "Registre parecer claro e objetivo.",
        "Confirme atualização para publicação no BI.",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Perfil Gestão (Acompanhamento BI)", level=2)
    for step in [
        "Abra o dashboard conectado à lista Agenda_Solicitacoes.",
        "Monitore KPIs: total de pedidos, taxa de aprovação, conflitos e lead time de decisão.",
        "Use filtros por equipe, período, tipo de pedido e status.",
    ]:
        doc.add_paragraph(step, style="List Number")


def add_operations_and_support(doc: Document) -> None:
    doc.add_heading("7. Operação Diária e Governança", level=1)
    checklist = [
        "Conferir fila Pendentes no início do turno.",
        "Priorizar pedidos Emergenciais.",
        "Garantir parecer em todas as recusas/reagendamentos.",
        "Revisar pedidos com conflito no fechamento do dia.",
        "Validar atualização do dataset no Power BI.",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8. Troubleshooting Rápido", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Sintoma"
    table.rows[0].cells[1].text = "Causa provável"
    table.rows[0].cells[2].text = "Ação recomendada"

    data = [
        ("Não salva data/hora", "Campo obrigatório vazio ou formato inválido", "Validar Inicio/Fim e campo Título/IDPedido"),
        ("Status não muda", "Fluxo pausado/erro de conexão", "Revisar histórico do Power Automate"),
        ("Conflito não marcado", "Filtro de sobreposição incorreto", "Validar condição Inicio_existente < Fim_novo e Fim_existente > Inicio_novo"),
        ("Parecer ausente", "Regra de obrigatoriedade não aplicada", "Forçar validação no Power Apps e no fluxo"),
    ]
    for row in data:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[2]

    doc.add_heading("9. Segurança e Perfis", level=1)
    doc.add_paragraph(
        "Usar grupo de segurança (Entra ID) para administradores. Evitar senha compartilhada de ADM. "
        "Trabalhar com contas nominativas para auditoria, rastreabilidade e conformidade."
    )

    doc.add_heading("10. Critério de Pronto para Produção", level=1)
    ready = [
        "Fluxos A, B, C e D testados com evidência.",
        "Cores de status funcionando conforme regra.",
        "Bloqueio de repetição em mesmo horário rejeitado validado.",
        "Dashboard BI publicado com filtros e KPIs.",
        "Treinamento rápido dos perfis concluído.",
    ]
    for item in ready:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()
    set_base_styles(doc)
    add_title(doc)
    add_intro(doc)
    add_list_config(doc)
    add_automate_config(doc)
    add_powerapps_config(doc)
    add_user_manual(doc)
    add_operations_and_support(doc)
    doc.save(str(OUT_DOCX))
    print(str(OUT_DOCX))


if __name__ == "__main__":
    main()
